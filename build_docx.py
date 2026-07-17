from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import os

# Files and paths
IMG_NAME = 'figure.png'
DOCX_NAME = 'manuscript.docx'

# Manuscript content
title = 'Computational Analysis of Von Neumann Entropy in One-Dimensional Spin Chains with Golden Ratio Geometric Coupling'
author = 'Fabian Leo Naressi'
affil = 'Independent Researcher'
date = 'July 2026'

abstract = (
    "This study investigates the influence of geometric coupling profiles on bipartite entanglement in one-dimensional spin chains. "
    "We analyze a transverse-field Ising model in which nearest-neighbor interactions are assigned according to three different spatial patterns: constant coupling, linear coupling, and a geometric decay governed by the Golden Ratio, φ. "
    "The ground state of the system is obtained through exact diagonalization using the QuTiP numerical framework, and the degree of entanglement is quantified by the Von Neumann entropy of a reduced density matrix corresponding to one half of the chain. "
    "The results indicate that the Golden Ratio-based coupling profile produces a distinctive entropy response near the scaling parameter α ≈ φ, yielding a broader and more pronounced entanglement profile than the constant and linear cases. These observations suggest that nontrivial geometric organization of interactions may enhance the persistence of quantum correlations in low-dimensional many-body systems."
)

methods = (
    "Model and Methods:\n"
    "We consider a transverse-field Ising chain of N spin-1/2 particles. The site-dependent coupling used in this study follows:\n"
    "J_n = alpha * phi^(-n) * exp(-n/phi^2), where phi is the Golden Ratio and alpha the global scale. "
    "Substituting this into the Hamiltonian gives:"
    "\nH = - sum_{n=0}^{N-2} (alpha phi^{-n} e^{-n/phi^2}) sigma_x^{(n)} sigma_x^{(n+1)} - g sum_{n=0}^{N-1} sigma_z^{(n)}.\n"
)

results = (
    "Results and Discussion:\n"
    "The geometric (Golden-Ratio) coupling suppresses early entropy growth and produces a stable inflection near alpha = phi. Energetically, the geometric chain stabilizes at a systematically higher ground-state energy (less negative) than the linear model, indicating a more optimized configuration that reduces overall binding energy required for global coherence."
)

fig_caption = (
    "Figure 1. Entanglement and ground-state energy comparison between linear and Golden-Ratio geometric couplings. "
    "(Left) Von Neumann entanglement entropy S_VN (bits) of the left-half reduced density matrix versus the scaling parameter alpha. "
    "The grey dashed curve shows the classical linear-cable model; the solid blue curve shows Naressi’s geometric law J_n = alpha phi^{-n} e^{-n/phi^2}. "
    "The vertical red dashed line marks the critical resonance alpha = phi ≈ 1.618. (Right) Ground-state energy E_0 (arbitrary units) as a function of alpha: the geometric law (solid green) remains systematically less negative than the linear model, indicating a shallower binding energy and a more energetically optimal stabilization of the correlated ground state."
)

# Create placeholder image if real figure not present
if not os.path.exists(IMG_NAME):
    W, H = 1600, 600
    img = Image.new('RGB', (W, H), color=(255,255,255))
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype('arial.ttf', 28)
    except Exception:
        font = ImageFont.load_default()
    text = 'Placeholder for figure: replace with high-resolution figure.png'
    try:
        bbox = d.textbbox((0, 0), text, font=font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
    except Exception:
        try:
            tw, th = font.getsize(text)
        except Exception:
            tw, th = (len(text) * 6, 20)
    d.text(((W-tw)/2, (H-th)/2), text, fill=(0,0,0), font=font)
    img.save(IMG_NAME)

# Build the Word document
doc = Document()

# Title page
doc.add_heading(title, level=0)
doc.add_paragraph(f'{author}\n{affil}\n{date}')
doc.add_page_break()

# Abstract
doc.add_heading('Abstract', level=1)
doc.add_paragraph(abstract)

# Keywords
doc.add_heading('Keywords', level=2)
doc.add_paragraph('Quantum entanglement; Von Neumann entropy; spin chains; Golden Ratio; transverse-field Ising model; QuTiP')

# Model and Methods
doc.add_heading('1. Introduction and Hamiltonian Model', level=1)
doc.add_paragraph(methods)

# Results
doc.add_heading('2. Results and Discussion', level=1)
doc.add_paragraph(results)

# Insert figure
doc.add_heading('Figure 1', level=2)
doc.add_paragraph(fig_caption)
# add image scaled to page width
try:
    doc.add_picture(IMG_NAME, width=Inches(6.5))
except Exception as e:
    print('Warning: could not insert image:', e)

# Acknowledgments
doc.add_heading('Acknowledgments', level=1)
doc.add_paragraph('The author thanks the developers of QuTiP and the open-source community for computational tools supporting this work.')

# References
doc.add_heading('References', level=1)
doc.add_paragraph(
    'Amico, L., Fazio, R., Osterloh, A., & Vedral, V. (2008). Entanglement in many-body systems. Reviews of Modern Physics, 80(2), 517–576.\n'
    'Johansson, J. R., Nation, P. D., & Nori, F. (2013). QuTiP 2: A Python framework for the dynamics of open quantum systems. Computer Physics Communications, 184(4), 1234–1240.\n'
    'Preskill, J. (2018). Quantum computing in the NISQ era and beyond. Quantum, 2, Article 79.\n'
)

# Save
doc.save(DOCX_NAME)
print('Created', DOCX_NAME)
